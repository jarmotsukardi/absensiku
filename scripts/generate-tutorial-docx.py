#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SHOT_DIR = ROOT / "public" / "tutorials" / "screenshots"
OUTPUT_DOCX = ROOT / "public" / "tutorials" / "tutorial-absensiku-admin-pegawai.docx"


def add_heading(document: Document, text: str, level: int = 1):
  document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str):
  p = document.add_paragraph(text)
  p.paragraph_format.space_after = Pt(8)


def add_image(document: Document, filename: str, caption: str, mobile: bool = False):
  path = SHOT_DIR / filename
  if not path.exists():
    add_paragraph(document, f"[Gambar tidak ditemukan: {filename}]")
    return

  width = Inches(2.6) if mobile else Inches(6.3)
  img = document.add_picture(str(path), width=width)
  last = document.paragraphs[-1]
  last.alignment = WD_ALIGN_PARAGRAPH.CENTER
  cap = document.add_paragraph(caption)
  cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
  cap.runs[0].italic = True
  cap.runs[0].font.size = Pt(9)


def main():
  doc = Document()
  title = doc.add_heading("Tutorial AbsensiKu: Admin Organisasi dan Pegawai", level=0)
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER

  add_paragraph(
    doc,
    "Dokumen ini memandu alur utama penggunaan AbsensiKu dari registrasi admin organisasi, setup absensi, operasional harian admin, hingga penggunaan oleh pegawai.",
  )
  add_paragraph(
    doc,
    "Catatan: Absensi tetap menjadi fondasi utama produk. Modul HR dan Payroll berada di dalam aplikasi sebagai jalur lanjutan setelah organisasi stabil pada operasional absensi inti. Tutorial ini tetap fokus pada fondasi absensi agar onboarding awal lebih cepat dan jelas.",
  )

  add_heading(doc, "Prasyarat", level=1)
  for item in [
    "URL aplikasi aktif, contoh: http://127.0.0.1:5173.",
    "Admin organisasi memiliki email aktif.",
    "Data master minimum disiapkan (OPD/unit/lokasi kerja sesuai kebutuhan).",
    "Pegawai menerima akun (via undangan/admin) untuk login.",
  ]:
    doc.add_paragraph(item, style="List Number")

  add_heading(doc, "A. Tutorial Admin Organisasi", level=1)
  admin_steps = [
    (
      "1) Buka Login Admin Organisasi",
      "Akses halaman /org/login. Tab Masuk digunakan untuk admin yang sudah punya akun.",
      "01-org-login.png",
      False,
      "Gambar 01 - Login Admin Organisasi",
    ),
    (
      "2) Daftar Organisasi (Admin Baru)",
      "Pilih tab Daftar Organisasi, isi data organisasi, lalu submit. Setelah berhasil, lanjut login sebagai admin organisasi.",
      "02-org-register.png",
      False,
      "Gambar 02 - Daftar Organisasi",
    ),
    (
      "3) Masuk ke Dashboard Organisasi",
      "Login menggunakan email dan password admin organisasi. Sistem akan mengarahkan ke dashboard utama organisasi.",
      "03-org-dashboard.png",
      False,
      "Gambar 03 - Dashboard Organisasi",
    ),
    (
      "4) Setup Awal Absensi",
      "Buka /org/onboarding lalu lengkapi konfigurasi dasar tenant dan preferensi operasional.",
      "04-org-onboarding.png",
      False,
      "Gambar 04 - Onboarding Organisasi",
    ),
    (
      "5) Atur Jam Kerja",
      "Buka /org/schedule/work-hours. Atur jam masuk, jam pulang, toleransi keterlambatan, dan aturan shift.",
      "05-org-work-hours.png",
      False,
      "Gambar 05 - Pengaturan Jam Kerja",
    ),
    (
      "6) Kelola Data Pegawai Aktif",
      "Buka /org/employees/active untuk tambah/edit/nonaktifkan data pegawai sesuai kebutuhan.",
      "06-org-employees-active.png",
      False,
      "Gambar 06 - Pegawai Aktif",
    ),
    (
      "7) Kirim Undangan Pegawai",
      "Buka /org/invitations. Buat undangan pegawai dan pantau status undangan.",
      "07-org-invitations.png",
      False,
      "Gambar 07 - Undangan Pegawai",
    ),
    (
      "8) Kelola Permohonan Kehadiran",
      "Buka /org/leave/requests. Setujui atau tolak permohonan dengan catatan yang jelas.",
      "08-org-leave-requests.png",
      False,
      "Gambar 08 - Permohonan Kehadiran",
    ),
    (
      "9) Laporan Absensi",
      "Buka /org/reports/attendance untuk memantau rekap absensi dan kebutuhan evaluasi periodik.",
      "09-org-report-attendance.png",
      False,
      "Gambar 09 - Laporan Absensi",
    ),
    (
      "10) FAQ Organisasi",
      "Buka /org/help/faq untuk melihat panduan operasional dan troubleshooting.",
      "10-org-help-faq.png",
      False,
      "Gambar 10 - FAQ Organisasi",
    ),
    (
      "11) Buat Tiket Bantuan",
      "Buka /org/help/tickets untuk membuat tiket kendala dan memantau progres penyelesaiannya.",
      "11-org-help-ticket.png",
      False,
      "Gambar 11 - Tiket Bantuan Organisasi",
    ),
  ]

  for heading, desc, image, mobile, caption in admin_steps:
    add_heading(doc, heading, level=2)
    add_paragraph(doc, desc)
    add_image(doc, image, caption, mobile=mobile)

  add_heading(doc, "B. Tutorial Pegawai", level=1)
  employee_steps = [
    (
      "1) Login Pegawai",
      "Buka /employee/login lalu masuk menggunakan akun pegawai aktif.",
      "12-employee-login-mobile.png",
      True,
      "Gambar 12 - Login Pegawai (Mobile)",
    ),
    (
      "2) Registrasi Pegawai",
      "Pada tab Daftar, gunakan mode Email atau Undangan sesuai kebijakan organisasi.",
      "13-employee-register-mobile.png",
      True,
      "Gambar 13 - Registrasi Pegawai (Mobile)",
    ),
    (
      "3) Dashboard Pegawai",
      "Setelah login, pegawai masuk ke dashboard utama untuk absensi dan fitur operasional harian.",
      "14-employee-dashboard-mobile.png",
      True,
      "Gambar 14 - Dashboard Pegawai (Mobile)",
    ),
    (
      "4) Ajukan Permohonan",
      "Buka /employee/dashboard?tab=requests untuk mengirim pengajuan seperti izin/cuti/izin terlambat/pulang cepat.",
      "15-employee-requests-mobile.png",
      True,
      "Gambar 15 - Pengajuan Pegawai (Mobile)",
    ),
    (
      "5) Riwayat Absensi",
      "Buka /employee/dashboard?tab=history untuk melihat status absensi harian.",
      "16-employee-history-mobile.png",
      True,
      "Gambar 16 - Riwayat Pegawai (Mobile)",
    ),
    (
      "6) Bantuan Pegawai",
      "Buka /employee/dashboard?tab=help untuk FAQ dan panduan penggunaan aplikasi.",
      "17-employee-help-mobile.png",
      True,
      "Gambar 17 - Bantuan Pegawai (Mobile)",
    ),
    (
      "7) Profil Pegawai",
      "Buka /employee/profile untuk melihat dan memperbarui data profil yang diizinkan.",
      "18-employee-profile-mobile.png",
      True,
      "Gambar 18 - Profil Pegawai (Mobile)",
    ),
  ]

  for heading, desc, image, mobile, caption in employee_steps:
    add_heading(doc, heading, level=2)
    add_paragraph(doc, desc)
    add_image(doc, image, caption, mobile=mobile)

  add_heading(doc, "C. Checklist Go-Live", level=1)
  for item in [
    "Jam kerja dan aturan absensi sudah ditetapkan.",
    "Lokasi kerja/radius sudah diverifikasi.",
    "Data pegawai aktif bersih dari duplikasi akun.",
    "Undangan pegawai terkirim dan teraktivasi.",
    "Alur approval permohonan berjalan konsisten.",
    "Kanal bantuan (FAQ + tiket) sudah disosialisasikan ke pengguna.",
  ]:
    doc.add_paragraph(item, style="List Number")

  add_heading(doc, "D. Lokasi File Tutorial", level=1)
  add_paragraph(doc, "DOCX editable: /tutorials/tutorial-absensiku-admin-pegawai.docx")
  add_paragraph(doc, "Versi web: /tutorials/tutorial-absensiku-admin-pegawai.html")
  add_paragraph(doc, "Folder screenshot: /tutorials/screenshots/")

  OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
  doc.save(str(OUTPUT_DOCX))
  print(f"[ok] generated {OUTPUT_DOCX}")


if __name__ == "__main__":
  main()
